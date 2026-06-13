# ===
# File Summary
# Path: backend\tests\documents\test_extraction.py
# Type: python
# Purpose: Backend test suite validating domain invariants and API behavior.
# Primary responsibilities:
# - Domain behavior is summarized for fast onboarding and avoids full-file reread.
# - Core symbols: test_docx_extraction_reads_paragraphs_and_table_cells, test_docx_extraction_caps_extracted_text, test_xlsx_extraction_reads_visible_sheets_only, test_pdf_extraction_reads_text_layer, test_image_ocr_gracefully_handles_missing_tesseract
# Inputs:
# - Downstream and upstream interactions in the same domain.
# Outputs:
# - API payloads, records, side effects, or UI views depending on file role.
# Dependencies:
# - Shared runtime services and adjacent domain modules.
# Known risks:
# - Validate behavior after migrations, dependency upgrades, or contract changes.
# ===
# 

import pytest

from apps.documents.extraction import MAX_EXTRACTED_TEXT_CHARS, _sanitize_text, extract_text


def test_docx_extraction_reads_paragraphs_and_table_cells(tmp_path):
    from docx import Document

    path = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("Material specification")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Grade"
    table.cell(0, 1).text = "PE-100"
    document.save(path)

    text, status = extract_text(
        path,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        path.name,
    )

    assert status == "extracted"
    assert "Material specification" in text
    assert "Grade" in text
    assert "PE-100" in text


def test_docx_extraction_caps_extracted_text(tmp_path):
    from docx import Document

    path = tmp_path / "large.docx"
    document = Document()
    document.add_paragraph("x" * (MAX_EXTRACTED_TEXT_CHARS + 1000))
    document.save(path)

    text, status = extract_text(
        path,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        path.name,
    )

    assert status == "extracted"
    assert len(text) == MAX_EXTRACTED_TEXT_CHARS


def test_extraction_removes_nul_bytes_before_database_save():
    text = _sanitize_text("polycarbonate\x00tensile\x00data")

    assert "\x00" not in text
    assert text == "polycarbonatetensiledata"


def test_xlsx_extraction_reads_visible_sheets_only(tmp_path):
    from openpyxl import Workbook

    path = tmp_path / "sample.xlsx"
    workbook = Workbook()
    visible = workbook.active
    visible.title = "Visible"
    visible["A1"] = "Lot"
    visible["B1"] = "L-001"
    hidden = workbook.create_sheet("Hidden")
    hidden.sheet_state = "hidden"
    hidden["A1"] = "Do not index"
    workbook.save(path)

    text, status = extract_text(
        path,
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        path.name,
    )

    assert status == "extracted"
    assert "Visible" in text
    assert "Lot" in text
    assert "L-001" in text
    assert "Do not index" not in text


def test_pdf_extraction_reads_text_layer(tmp_path):
    from pypdf import PdfReader, PdfWriter
    from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

    path = tmp_path / "sample.pdf"
    writer = PdfWriter()
    page = writer.add_blank_page(width=300, height=144)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    page[NameObject("/Resources")] = DictionaryObject(
        {
            NameObject("/Font"): DictionaryObject({NameObject("/F1"): writer._add_object(font)})
        }
    )
    content = DecodedStreamObject()
    content.set_data(b"BT /F1 18 Tf 20 100 Td (Searchable polymer PDF) Tj ET")
    page[NameObject("/Contents")] = writer._add_object(content)
    with path.open("wb") as handle:
        writer.write(handle)
    assert "Searchable polymer PDF" in (PdfReader(path).pages[0].extract_text() or "")

    text, status = extract_text(path, "application/pdf", path.name)

    assert status == "extracted"
    assert "Searchable polymer PDF" in text


def test_image_ocr_gracefully_handles_missing_tesseract(tmp_path, monkeypatch):
    from PIL import Image

    path = tmp_path / "sample.png"
    image = Image.new("RGB", (32, 32), color="white")
    image.save(path)

    try:
        import pytesseract
    except ImportError:
        pytesseract = None

    if pytesseract is not None:
        monkeypatch.setattr(
            pytesseract,
            "image_to_string",
            lambda _image: (_ for _ in ()).throw(RuntimeError("tesseract missing")),
        )

    text, status = extract_text(path, "image/png", path.name)

    assert text == ""
    assert status in {"unsupported", "failed"}


@pytest.mark.parametrize("file_name,mime_type", [("notes.txt", "text/plain"), ("archive.bin", "")])
def test_unknown_file_types_are_unsupported(tmp_path, file_name, mime_type):
    path = tmp_path / file_name
    path.write_bytes(b"plain text is not an indexed document format")

    text, status = extract_text(path, mime_type, file_name)

    assert text == ""
    assert status == "unsupported"

