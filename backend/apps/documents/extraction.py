# ===
# File Summary
# Path: backend\apps\documents\extraction.py
# Type: python
# Purpose: Document domain service managing records, revisions, and extraction workflows.
# Primary responsibilities:
# - Domain behavior is summarized for fast onboarding and avoids full-file reread.
# - Core symbols: extract_text, _extract_pdf, _extract_docx, _extract_xlsx, _extract_image
# Inputs:
# - Downstream and upstream interactions in the same domain.
# Outputs:
# - API payloads, records, side effects, or UI views depending on file role.
# Dependencies:
# - Shared runtime services and adjacent domain modules.
# Known risks:
# - Validate behavior after migrations, dependency upgrades, or contract changes.
# ===
# 

from pathlib import Path


MAX_EXTRACTED_TEXT_CHARS = 200_000
MAX_PDF_PAGES = 50
MAX_XLSX_ROWS_PER_SHEET = 5_000
MAX_XLSX_TOTAL_CELLS = 50_000
OCR_TIMEOUT_SECONDS = 10

PDF_MIME_TYPES = {"application/pdf"}
DOCX_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
XLSX_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}


def extract_text(path, mime_type: str, file_name: str) -> tuple[str, str]:
    path = Path(path)
    mime_type = (mime_type or "").lower()
    extension = path.suffix.lower() or Path(file_name).suffix.lower()

    if mime_type in PDF_MIME_TYPES or extension == ".pdf":
        return _extract_pdf(path)
    if mime_type in DOCX_MIME_TYPES or extension == ".docx":
        return _extract_docx(path)
    if mime_type in XLSX_MIME_TYPES or extension == ".xlsx":
        return _extract_xlsx(path)
    if mime_type.startswith("image/") or extension in {".png", ".jpg", ".jpeg", ".tif", ".tiff"}:
        return _extract_image(path)
    return "", "unsupported"


def _extract_pdf(path: Path) -> tuple[str, str]:
    try:
        from pypdf import PdfReader

        reader = PdfReader(path)
        parts = []
        for page in reader.pages[:MAX_PDF_PAGES]:
            parts.append(page.extract_text() or "")
            if _joined_length(parts) >= MAX_EXTRACTED_TEXT_CHARS:
                break
        text = _cap_text("\n".join(filter(None, parts)))
    except Exception:
        return "", "failed"
    return text, "extracted" if text.strip() else "unsupported"


def _extract_docx(path: Path) -> tuple[str, str]:
    try:
        from docx import Document

        document = Document(path)
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
        for table in document.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells if cell.text)
        text = _cap_text("\n".join(parts))
    except Exception:
        return "", "failed"
    return text, "extracted" if text.strip() else "unsupported"


def _extract_xlsx(path: Path) -> tuple[str, str]:
    try:
        from openpyxl import load_workbook

        workbook = load_workbook(path, read_only=True, data_only=True)
        parts = []
        total_cells = 0
        for sheet in workbook.worksheets:
            if sheet.sheet_state != "visible":
                continue
            parts.append(sheet.title)
            for row_index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
                if row_index > MAX_XLSX_ROWS_PER_SHEET or total_cells >= MAX_XLSX_TOTAL_CELLS:
                    break
                values = [str(value) for value in row if value is not None and str(value) != ""]
                total_cells += len(row)
                if values:
                    parts.append("\t".join(values))
                if _joined_length(parts) >= MAX_EXTRACTED_TEXT_CHARS:
                    break
            if total_cells >= MAX_XLSX_TOTAL_CELLS or _joined_length(parts) >= MAX_EXTRACTED_TEXT_CHARS:
                break
        workbook.close()
        text = _cap_text("\n".join(parts))
    except Exception:
        return "", "failed"
    return text, "extracted" if text.strip() else "unsupported"


def _extract_image(path: Path) -> tuple[str, str]:
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        return "", "unsupported"

    try:
        with Image.open(path) as image:
            text = pytesseract.image_to_string(image, timeout=OCR_TIMEOUT_SECONDS)
    except Exception:
        return "", "failed"
    text = _cap_text(text)
    return text, "extracted" if text.strip() else "unsupported"


def _cap_text(text: str) -> str:
    return text[:MAX_EXTRACTED_TEXT_CHARS]


def _joined_length(parts) -> int:
    return sum(len(part) for part in parts) + max(len(parts) - 1, 0)

