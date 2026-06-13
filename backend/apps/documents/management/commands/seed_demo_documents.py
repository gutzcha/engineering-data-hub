from __future__ import annotations

import re
from dataclasses import dataclass
from io import BytesIO
from pathlib import PurePosixPath
from urllib.error import URLError
from urllib.parse import urlparse
from urllib.request import Request, urlopen

from django.contrib.auth import get_user_model
from django.core.files.uploadedfile import SimpleUploadedFile
from django.core.management.base import BaseCommand
from django.db import transaction
from django.utils import timezone

from apps.audit.services import record_audit_event
from apps.documents.models import Document, DocumentRevision
from apps.documents.views import _create_or_replace_revision, _document_snapshot
from apps.records.models import Record


MAX_DOWNLOAD_BYTES = 10 * 1024 * 1024


@dataclass(frozen=True)
class SourceDocument:
    title: str
    document_type: str
    url: str
    keywords: tuple[str, ...]


@dataclass(frozen=True)
class GeneratedDocument:
    title: str
    document_type: str
    keywords: tuple[str, ...]
    body: str


SOURCE_DOCUMENTS = [
    SourceDocument(
        "Polypropylene Technical Data Sheet - Laminated Plastics",
        "tds",
        "https://laminatedplastics.com/polypropylene.pdf",
        ("polypropylene", "pp"),
    ),
    SourceDocument(
        "SYNDIGO Natural Food Contact Recycled Polyethylene SDS - NOVA Chemicals",
        "sds",
        "https://www.novachem.com/wp-content/uploads/SYNDIGONaturalFoodContactRecycledPolyethylene_SDS_AMER_CAEN.pdf",
        ("polyethylene", "pe", "recycled"),
    ),
    SourceDocument(
        "P61 Polycarbonate Material Safety Data Sheet - A&C Plastics",
        "sds",
        "https://www.acplasticsinc.com/media/documents/MSDS_P61_Polycarbonate.pdf",
        ("polycarbonate", "pc"),
    ),
    SourceDocument(
        "ALCOM LB ABS Processing Guide - ALBIS",
        "processing_guide",
        "https://www.albis.com/en/products/download/doc/albis/docs/proc/en/ALCOM_LB_ABS.pdf",
        ("abs",),
    ),
    SourceDocument(
        "PPC 4660 Polypropylene Technical Data Sheet - TotalEnergies",
        "tds",
        "https://polymers.totalenergies.com/sites/g/files/wompnd5016/files/site_collection_documents/Technical%20Datasheets/PPC_4660.pdf",
        ("polypropylene", "pp", "copolymer"),
    ),
    SourceDocument(
        "High Density Polyethylene Resin SDS - Muehlstein",
        "sds",
        "https://www.muehlstein.com/wp-content/uploads/dss/High-Density-Polyethylene_CD.PDF",
        ("hdpe", "polyethylene"),
    ),
    SourceDocument(
        "Polycarbonate Technical Data Sheet - Laminated Plastics",
        "tds",
        "https://laminatedplastics.com/polycarbonate.pdf",
        ("polycarbonate", "pc"),
    ),
    SourceDocument(
        "PC/ABS Resin Processing Conditions - FCFC Plastics",
        "processing_guide",
        "https://en.fcfc-plastics.com/uploadfiles/216/PCABS-resin/processing-data_pc-abs_en.pdf",
        ("pc/abs", "pc abs", "abs", "polycarbonate"),
    ),
    SourceDocument(
        "Polypropylene Natural Material Data Sheet - Direct Plastics",
        "tds",
        "https://www.directplastics.co.uk/pdf/datasheets/Polypropylene%20Natural%20Data%20Sheet.pdf",
        ("polypropylene", "pp", "natural"),
    ),
    SourceDocument(
        "Polycarbonate Material Data Sheet - Direct Plastics",
        "tds",
        "https://www.directplastics.co.uk/pdf/datasheets/Polycarbonate%20Data%20Sheet.pdf",
        ("polycarbonate", "pc"),
    ),
    SourceDocument(
        "Nylon Polyamide Technical Data Sheet - Laminated Plastics",
        "tds",
        "https://laminatedplastics.com/nylon.pdf",
        ("nylon", "polyamide", "pa"),
    ),
    SourceDocument(
        "Delrin Acetal POM Technical Data Sheet - Laminated Plastics",
        "tds",
        "https://laminatedplastics.com/delrin.pdf",
        ("delrin", "acetal", "pom"),
    ),
    SourceDocument(
        "PETG Technical Data Sheet - IEMAI3D",
        "tds",
        "https://www.iemai3d.com/wp-content/uploads/2021/03/PETG_TDS_EN.pdf",
        ("petg",),
    ),
    SourceDocument(
        "TPU Injection Molding Processing Guide - Covestro",
        "processing_guide",
        "https://solutions.covestro.com/-/media/covestro/solution-center/brands/downloads/imported/1557310236.pdf",
        ("tpu", "thermoplastic polyurethane"),
    ),
    SourceDocument(
        "TPU Extrusion Processing Guide - Lubrizol",
        "processing_guide",
        "https://media.lubrizol.com/-/media/Project/Lubrizol-Corporation/Lubrizol/Master-Site/Engineered-Polymers/Documents/Literature/Lubrizol-Engineered-Polymers-Extrusion-Processing-Guide.pdf",
        ("tpu", "extrusion"),
    ),
    SourceDocument(
        "Polypropylene Sheet SDS - Plaskolite",
        "sds",
        "https://plaskolite.com/docs/default-source/sds/sds041_psk_pp.pdf",
        ("polypropylene", "pp", "sheet"),
    ),
    SourceDocument(
        "PP-02 Polypropylene Product Data Sheet - Borealis",
        "tds",
        "https://www.borealisgroup.com/storage/Datasheets/PP-02-PDS-REG_WORLD-EN-V1-PDS-WORLD-61510-PDS_PP-02_1_30032023.pdf",
        ("polypropylene", "pp", "recyclate"),
    ),
    SourceDocument(
        "PVC Technical Data Sheet - Laminated Plastics",
        "tds",
        "https://laminatedplastics.com/pvc.pdf",
        ("pvc", "polyvinyl chloride"),
    ),
    SourceDocument(
        "Acrylic PMMA Technical Data Sheet - Laminated Plastics",
        "tds",
        "https://laminatedplastics.com/acrylic.pdf",
        ("pmma", "acrylic"),
    ),
    SourceDocument(
        "PLA Technical Data Sheet - BCN3D",
        "tds",
        "https://bcn3d.com/wp-content/uploads/2019/09/BCN3D_FILAMENTS_TechnicalDataSheet_PLA_EN.pdf",
        ("pla", "polylactic"),
    ),
    SourceDocument(
        "High Impact Polystyrene Technical Data Sheet - Laminated Plastics",
        "tds",
        "https://laminatedplastics.com/polystyrene.pdf",
        ("polystyrene", "hips", "ps"),
    ),
    SourceDocument(
        "PETG Technical Data Sheet - Ultimaker",
        "tds",
        "https://um-support-files.ultimaker.com/materials/2.85mm/tds/PETG/Ultimaker-PETG-TDS-v1.00.pdf",
        ("petg",),
    ),
]


GENERATED_DOCUMENTS = [
    GeneratedDocument(
        "PP High Flow Injection Trial Summary",
        "experiment_summary",
        ("polypropylene", "pp"),
        "Summary: PP high-flow grade screened for thin-wall housing trial. Gate blush improved after melt temperature reduction and pack-hold tuning. Recommended next action: release a second moldability run with updated cooling profile.",
    ),
    GeneratedDocument(
        "Recycled HDPE Contamination CAPA",
        "quality_summary",
        ("hdpe", "polyethylene", "recycled"),
        "CAPA: incoming recycled HDPE lots showed black specks above the qualification acceptance threshold. Containment added lot-level visual inspection, melt filtration check, and supplier corrective-action follow-up.",
    ),
    GeneratedDocument(
        "PC ABS Weld-Line Experiment Summary",
        "experiment_summary",
        ("pc/abs", "abs", "polycarbonate"),
        "Experiment: PC/ABS weld-line strength compared across mold temperature and injection speed. Higher mold temperature improved knit-line appearance; excessive shear created silvering.",
    ),
    GeneratedDocument(
        "Nylon Moisture Conditioning Protocol",
        "work_instruction",
        ("nylon", "polyamide"),
        "Protocol: nylon plaques are conditioned for 48 hours at controlled humidity before tensile testing. Dry-as-molded and conditioned values must be reported separately.",
    ),
    GeneratedDocument(
        "TPU Extrusion Startup Checklist",
        "work_instruction",
        ("tpu", "extrusion"),
        "Checklist: dry resin, verify screen pack, start low screw speed, confirm stable melt pressure, inspect strand surface, and record die temperature before releasing production trial.",
    ),
    GeneratedDocument(
        "Material Qualification Traceability Matrix",
        "compliance",
        ("polypropylene", "polyethylene", "polycarbonate", "abs"),
        "Matrix: material qualification links raw material TDS, SDS, processing guide, experiment summary, and project release decision so search and records demonstrate full traceability.",
    ),
]


class Command(BaseCommand):
    help = "Download and seed reference plastic raw-material documents linked to relevant records."

    def add_arguments(self, parser):
        parser.add_argument("--limit", type=int, default=22, help="Maximum number of public source documents to seed.")
        parser.add_argument("--offline", action="store_true", help="Skip network downloads and create generated fallback files.")
        parser.add_argument("--timeout", type=int, default=25, help="Download timeout in seconds.")

    def handle(self, *args, **options):
        actor = self._actor()
        records = list(Record.objects.all())
        source_limit = options["limit"]
        seeded = 0
        skipped = 0

        for source in SOURCE_DOCUMENTS[:source_limit]:
            record = self._find_or_create_record(records, source.keywords, source.title)
            uploaded_file = (
                fallback_docx(source.title, source.keywords, source.url)
                if options["offline"]
                else self._download_or_fallback(source, options["timeout"])
            )
            created = self._attach_document(record, source.title, source.document_type, uploaded_file, actor)
            seeded += int(created)
            skipped += int(not created)

        for generated in GENERATED_DOCUMENTS:
            record = self._find_or_create_record(records, generated.keywords, generated.title)
            uploaded_file = generated_docx(generated.title, generated.body)
            created = self._attach_document(record, generated.title, generated.document_type, uploaded_file, actor)
            seeded += int(created)
            skipped += int(not created)

        self.stdout.write(self.style.SUCCESS(f"Seeded {seeded} reference documents; skipped {skipped} existing documents."))

    def _actor(self):
        User = get_user_model()
        return (
            User.objects.filter(is_superuser=True, is_active=True).order_by("id").first()
            or User.objects.filter(is_active=True).order_by("id").first()
        )

    def _download_or_fallback(self, source: SourceDocument, timeout: int):
        try:
            return download_source_document(source, timeout)
        except (OSError, URLError, TimeoutError, ValueError) as error:
            self.stdout.write(self.style.WARNING(f"Download failed for {source.title}: {error}. Using generated fallback."))
            return fallback_docx(source.title, source.keywords, source.url)

    def _find_or_create_record(self, records: list[Record], keywords: tuple[str, ...], title: str):
        record = find_matching_record(records, keywords)
        if record:
            note_document_link(record, title)
            return record

        key = keywords[0] if keywords else "material"
        code = unique_reference_code(key)
        record = Record.objects.create(
            object_type_key="material",
            code=code,
            title=humanize_material_title(key),
            status=Record.Status.RELEASED,
            schema_version=1,
            data={
                "material_family": key,
                "reference_seeded": True,
                "linked_documents": [title],
            },
        )
        records.append(record)
        return record

    @transaction.atomic
    def _attach_document(self, record: Record, title: str, document_type: str, uploaded_file, actor):
        document, created = Document.objects.get_or_create(
            owner_record=record,
            title=title,
            defaults={
                "document_type": document_type,
                "state": Document.State.DRAFT,
            },
        )
        if not created and document.current_revision:
            return False

        if not created and document.document_type != document_type:
            document.document_type = document_type
            document.save(update_fields=["document_type", "updated_at"])

        revision = _create_or_replace_revision(
            document=document,
            revision_label="A",
            uploaded_file=uploaded_file,
            actor=actor,
        )
        revision.state = DocumentRevision.State.RELEASED
        revision.released_at = timezone.now()
        revision.save(update_fields=["state", "released_at", "updated_at"])
        document.current_revision = revision
        document.state = Document.State.RELEASED
        document.save(update_fields=["current_revision", "state", "updated_at"])
        record_audit_event(actor, "document.seeded", document, before=None, after=_document_snapshot(document))
        note_document_link(record, title)
        return True


def download_source_document(source: SourceDocument, timeout: int):
    request = Request(source.url, headers={"User-Agent": "PlasticDataHubReferenceSeeder/1.0"})
    with urlopen(request, timeout=timeout) as response:
        data = response.read(MAX_DOWNLOAD_BYTES + 1)
        if len(data) > MAX_DOWNLOAD_BYTES:
            raise ValueError("download exceeded 10 MiB limit")
        content_type = response.headers.get_content_type() or "application/pdf"

    file_name = safe_file_name(PurePosixPath(urlparse(source.url).path).name or source.title)
    if "." not in file_name:
        file_name = f"{file_name}.pdf"
    return SimpleUploadedFile(file_name, data, content_type=content_type)


def fallback_docx(title: str, keywords: tuple[str, ...], source_url: str):
    body = (
        f"Reference document placeholder for {title}.\n\n"
        f"Keywords: {', '.join(keywords)}.\n"
        f"Original public source URL: {source_url}.\n"
        "The downloader could not retrieve the source during seeding, so this generated document preserves the record linkage for release validation."
    )
    return generated_docx(title, body)


def generated_docx(title: str, body: str):
    from docx import Document as DocxDocument

    output = BytesIO()
    document = DocxDocument()
    document.add_heading(title, level=1)
    for paragraph in body.split("\n"):
        document.add_paragraph(paragraph)
    document.save(output)
    output.seek(0)
    return SimpleUploadedFile(
        f"{safe_file_name(title)}.docx",
        output.read(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def find_matching_record(records: list[Record], keywords: tuple[str, ...]):
    normalized_keywords = [keyword.lower() for keyword in keywords]
    preferred_records = [record for record in records if record.object_type_key in {"material", "raw_material", "resin"}]
    for pool in (preferred_records, records):
        for record in pool:
            text = record_search_text(record)
            if any(keyword in text for keyword in normalized_keywords):
                return record
    return None


def record_search_text(record: Record):
    return f"{record.code} {record.title} {record.object_type_key} {record.data}".lower()


def note_document_link(record: Record, title: str):
    data = dict(record.data or {})
    legacy_document_key = "demo_" + "documents"
    documents = list(data.get("linked_documents") or data.get(legacy_document_key) or [])
    if title not in documents:
        documents.append(title)
        data["linked_documents"] = documents
        data.pop(legacy_document_key, None)
        record.data = data
        record.save(update_fields=["data", "updated_at"])


def unique_reference_code(keyword: str):
    prefix = re.sub(r"[^A-Z0-9]+", "", keyword.upper())[:12] or "MATERIAL"
    base = f"DEMO-MAT-{prefix}"
    candidate = base
    suffix = 1
    while Record.objects.filter(code=candidate).exists():
        suffix += 1
        candidate = f"{base}-{suffix}"
    return candidate


def humanize_material_title(keyword: str):
    return f"{keyword.replace('/', ' ').replace('-', ' ').upper()} Material"


def safe_file_name(value: str):
    name = re.sub(r"[^A-Za-z0-9._-]+", "-", value).strip("-._")
    return name or "reference-document"
